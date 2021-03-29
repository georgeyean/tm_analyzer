#! /usr/bin/env python
'''
Author @ George Yean.

Install on Max or Linux:
pip3 install python-docx
pip3 install geograpy3
pip3 install sox
pip install symspellpy

Need python 3.7.5


This file includes the steps of preprocessing:

-split word docs into less than 150-words files (natural breaks)
-remove stop words
-remove numbers
-remove punctuation
-lowercase
-stemming
-
'''

import re, sys, os, pdb, time
from os import listdir
from os.path import isfile, join
import docx
import csv
import logging
import shutil
import nltk
from nltk import PorterStemmer
from nltk.corpus import stopwords
from treatyutils import convert_pic
duration = 1  # seconds
freq = 440  # Hz

import multiprocessing as mp
pool = mp.Pool(processes=mp.cpu_count())

logging.basicConfig(format='%(asctime)s,%(msecs)d %(levelname)-8s [%(filename)s:%(lineno)d] %(message)s',
    datefmt='%Y-%m-%d:%H:%M:%S',
    level=logging.DEBUG)

stemmer= PorterStemmer()
nltk.download('stopwords')

peace_dir_w = "Word docs_Peace"
peace_dir_w_clean = "Word docs_Peace_clean"
comm_dir_w = "Word docs_CommNav"
comm_dir_w_clean = "Word docs_CommNav_clean"
short_dir = "short_texts"
final_dir_clean = "Word final_clean"

'''
 Preprocessing word files
'''
def preprocessWord():
    logging.info("Processing Word files")
    
    if os.path.exists('short_texts'): 
        shutil.rmtree('short_texts')
    if os.path.exists('graph'):
        shutil.rmtree('graph')
    
    time.sleep(10) # wait 10s to ensure delete all
        
    os.mkdir('short_texts')
    os.mkdir('graph')
    
    peacefiles_w = [join(peace_dir_w_clean,f) for f in listdir(peace_dir_w_clean) if isfile(join(peace_dir_w_clean, f)) and ".docx" in f and "~$" not in f]
    logging.info("Peace treaties: %d"%(len(peacefiles_w)))

    commfiles_w = [join(comm_dir_w_clean,f) for f in listdir(comm_dir_w_clean) if isfile(join(comm_dir_w_clean, f)) and ".docx" in f and "~$" not in f]
    logging.info("Comm treaties: %d"%(len(commfiles_w)))
    
    files_w = peacefiles_w + commfiles_w
    files_w.sort()
    
    f = open('ShortBook.csv', 'w+')
    short_doc = csv.writer(f)
    short_doc.writerow(["fullBook", "shortFile", "fullFile"])
    
    logging.info("File amount: %d"%(len(files_w)))
    
    
    for f in range(len(files_w)): 
    
        #logging.info("Splitting: %s"%(files_w[f]))
        
        doc = docx.Document(files_w[f])
        result = [p.text for p in doc.paragraphs]

        #break document into sub-docs by natureal breaks. Merge until 150 words if too short
        file_name = join(short_dir, re.sub('.docx', '', files_w[f].split("/")[1]))
        logging.info("Filename: %s"%(file_name))
        
        merge_text = ""
        doc_num = 0
        
        for j in range(len(result)):
            try:
                if len(result[j].split()) > 150:
                    if len(merge_text.split()) > 0:
                        merge_text += " "+result[j]
                    else:
                        merge_text = result[j]
                    short_f = '%s_%d.txt'%(file_name, doc_num)
                    rfile = open(short_f, 'w')
                    rfile.write('%s'%(merge_text))
                    rfile.write('\n\n')
                    merge_text = ""
                    doc_num += 1
                    short_doc.writerow([files_w[f].split("/")[1], short_f.split("/")[1], short_f])
                else:
                    merge_text += " "+result[j]
                    if len(merge_text.split()) > 150:
                        short_f = '%s_%d.txt'%(file_name, doc_num)
                        rfile = open(short_f, 'w')
                        rfile.write('%s'%(merge_text))
                        rfile.write('\n\n')
                        merge_text = ""
                        doc_num += 1
                        short_doc.writerow([files_w[f].split("/")[1], short_f.split("/")[1], short_f])
            
            except Exception:
                logging.info("miss %s ", c['oup_reference'].strip(" "))
                continue         
                
    shortfile_t = [join(short_dir,f) for f in listdir(short_dir) if isfile(join(short_dir, f)) and ".txt" in f]
    logging.info("Short files total: %d"%(len(shortfile_t)))           

'''
Generate final keys for the model
'''
def genFinalKeys():
    logging.info("genFinalKeys")
    
    shortfile_t = [join(short_dir,f) for f in listdir(short_dir) if isfile(join(short_dir, f)) and ".txt" in f]
    
    dict = {}
    dict_appear = {}
    z = 0
    for i in range(len(shortfile_t)):
        z += 1
        logging.info("file %d"%z)
        f = open(shortfile_t[i], 'r', encoding="ISO-8859-1")
        used = {}
        text = ' '.join(f.readlines())
        #print("Lord" in text)
        text = text.lower()
        text = re.sub('- ', '', text)         
        text = re.sub('\W', ' ', text)
        text = re.sub('\d', ' ', text)

        splittext = text.split()
        splittext = map(lambda x: stemmer.stem(x), splittext)

        for word in splittext:
            used[word] = 1
            if word in dict:
                 dict[word] += 1
            else:
                 dict[word] = 1
    
        for k in used.keys():
            if k in dict_appear:
                dict_appear[k] += 1
            else:
                dict_appear[k] = 1
    
    logging.info("After steming: %d"%(len(dict)))
    logging.info("After steming appear: %d"%(len(dict_appear)))
    

    remove_thresh = []
    threshold = len(shortfile_t) * 0.005
    logging.info("Threshold: %f"%(threshold))
    
    for k in dict_appear.copy():
        if int(dict_appear[k]) < threshold:
            remove_thresh.append(k)
    logging.info("Remove threshold: %d"%(len(remove_thresh)))
    
    for k in list(dict):
        if k in remove_thresh:
            del dict[k]
    logging.info("After removed threshold: %d"%(len(dict)))    
    
    
    #remove stop words
    stops= stopwords.words('english')
    for s in range(len(stops)):
        stops[s] = stemmer.stem(stops[s])
        
    for k in list(dict):
        if k in stops:
            del dict[k]
    logging.info("After stop words remove: %d"%(len(dict)))   
    
    #manual remove
    man_remove = ["en", "e"]
    romans = ["i", "xi", "xxi", "ii", "xii", "xxii", "iii", "xiii", "xxiii", "iv", "xiv", "xxiv", \
    "v", "xv", "xxv", "vi", "xvi", "xxvi", "vii", "xvii", "xxvii", "viii", "xviii", "xxviii", "ix",\
    "xix", "xxix", "x", "xx", "xxx", "xxxiv", "xxxii"]
    for k in list(dict):
        if k in man_remove:
            del dict[k]
            continue
        if k in romans:
            del dict[k] 
            continue
        if len(k) < 2:
            del dict[k] 
            continue
    logging.info("After manaual remove: %d"%(len(dict)))       
    
    #remove less than 3 chars
    for k in list(dict):
        if len(k) < 3:
            del dict[k]
    logging.info("After shortchar remove: %d"%(len(dict)))    
    
    #remove geo
    geos = ["america", "spanish", "colombia", "brazil", "chile", "argentin", "costa", "rica", "dominican", "ecuador", "guatemala", "hondura", \
    "mexico", "nicaragua", "paraguay", "peru", "salvador", "uruguay", "venezuela", "derecho", "bangkok", "french", "india", "itali", "england",\
    "china", "chines", "sweden", "norway", "swedish", "austrian", "bolivia", "japan", "japanes",\
    "dutch", "german", "bolivian", "washington", "italian", "bavaria", "germani", "bulgaria", "hawaiian",\
    "poland", "africa", "african", "peruvian", "russia", "british", "english", "hungarian", "siam", "siames",\
    "austria", "american", "netherland", "portugues", "prussian", "prussia", "bulgarian", "norwegian",\
    "paraguayan","portug","belgium","russian","indian","belgian","brazilian","madrid", "britain", "sulina"]   
    for k in list(dict):
        if k in geos:
            del dict[k]
    logging.info("After geo remove: %d"%(len(dict)))               
    print(dict)
    
    #save keys
    f = open('ShortFinalKeys.txt', 'w')
    f.truncate(0)
    for k in list(dict):
        f.write(k)
        f.write('\n')
        

  
    f =  open('ShortBook.csv', 'r')
    starts = f.readlines()
    clean_starts = []
    
    for i in range(1, len(starts)):
        #ee = re.findall('truelaw|alchemy|consolationforruler', starts[z].split(',')[0].lower())
        #if len(ee)==0:
        clean_starts.append(starts[i].split(',')[-1].strip('\n'))
        
    # f = open('RevisedShortBooks.csv', 'w+')
    # new_files = csv.writer(f)
    # new_files.write(["fullBook","bookID","shortFile","fullFile"])
            
    #create termDoc
    logging.info("Creating doc-term file, please wait...")  
    t = open('ShortTermDoc.csv', 'w+')
    term_doc = csv.writer(t)

    # clean_text = ''
    # for word in list(dict):
    #     clean_text+= str(word)
    #     if word != list(dict)[-1]:
    #         clean_text+= ','
    
    term_doc.writerow(list(dict))
    #term_doc.write('\n')
    
    for j in range(len(clean_starts)):
        if j%1000 == 0:
            logging.info("Dont worry!!! scanning short file: %d %s" %(j, clean_starts[j])) 
        
        f = open(clean_starts[j], 'r', encoding="ISO-8859-1")
        text = ' '.join(f.readlines())
        #lord_cap = len(re.findall('Lord', text))
        #text = re.sub('Lord', ' ', text)
        text = text.lower()
        text = re.sub('- ', '', text)         
        text = re.sub('\W', ' ', text)
        text = re.sub('\d', ' ', text)
        splittext = text.split()
        splittext = map(lambda x: stemmer.stem(x), splittext)
        dict_dtm = {}
        for word in splittext:
            if word in dict_dtm:
                dict_dtm[word] += 1
    	    #used.append(word)
            else:
                dict_dtm[word] = 1
                # used.append(word)
                
        output = []
        for word in list(dict):
            if word in dict_dtm:
                output.append(str(dict_dtm[word]))
            if word not in dict_dtm:
                output.append(str(0))
            #if word != list(dict)[-1]:
            #    output += ','
        #output += str(lord_cap)       
        term_doc.writerow(output)
        if j in [1,5,8]:
            logging.info("DTM line %d: %s"%(j, output)) 
  
    
    
def genPeaceCommDoc():
    f =  open('alltreaties_metadata.csv', 'r')
    corrs = csv.DictReader(f)
    keys = {}
    for c in corrs:
        try:
            #print(c["date_force"].strip(" "))
            file = "_".join(c["oup_reference"].split(" "))
            year =  re.search(r"(1\d{3})", c["date_force"].strip(" ")).group(1)
            #print(year)
            keys[file] = int(year)
        except Exception:
            logging.info("miss %s ", c['oup_reference'].strip(" "))
            continue
            
    #print(keys)
    
    t = open('BooksPeaceComm.csv', 'w+')
    doc = csv.writer(t)
    doc.writerow(["Text", "Peace", "Year", "Abbreviation"])
    
    peacefiles_w = [f for f in listdir(peace_dir_w) if isfile(join(peace_dir_w, f)) and ".docx" in f and "~$" not in f]

    for f in range(len(peacefiles_w)):
        if any(c in ["$", "~"] for c in peacefiles_w[f]):
            print("quit ", peacefiles_w[f])
            continue
        file_sym = re.search(r"(\S+)_eng\S+", peacefiles_w[f]).group(1)
        #print(peacefiles_w[f])
        #print(file_sym)
        year = keys[file_sym]
        #print(file_sym, year)
        doc.writerow([peacefiles_w[f], "1", year, file_sym])

    commfiles_w = [f for f in listdir(comm_dir_w) if isfile(join(comm_dir_w, f)) and ".docx" in f and "~$" not in f]
    print(len(peacefiles_w), len(commfiles_w))
    for f in range(len(commfiles_w)):
        if any(c in ["$", "~"] for c in commfiles_w[f]):
            print("quit ", commfiles_w[f])
            continue
        file_sym = re.search(r"([\S]+)_eng\S+", commfiles_w[f]).group(1)
        #print(commfiles_w[f])
        #print(file_sym)
        year = keys[file_sym]
        #print(file_sym, year)
        doc.writerow([commfiles_w[f], "0", year, file_sym])
    


#Batch clean and create new docs based on corrective key/value set
#Note: check after finishing
#in peace/42_CTS_49_eng_text.docx 
#if mi- nisters is fixed
#if defence is changed to defense
def cleanTexts():
    
    if os.path.exists(peace_dir_w_clean):    
        shutil.rmtree(peace_dir_w_clean)
    if os.path.exists(comm_dir_w_clean):
        shutil.rmtree(comm_dir_w_clean)
    if os.path.exists(final_dir_clean):
        shutil.rmtree(final_dir_clean)
        
    os.mkdir(peace_dir_w_clean)
    os.mkdir(comm_dir_w_clean)
    os.mkdir(final_dir_clean)
    
    #pdb.set_trace()
    
    f =  open('spelling_corrections.csv', 'r')
    corrs = f.readlines()
    keys = {}
    for c in corrs:
        keys[c.split(',')[0].strip(" ")] = c.split(",")[1].strip(" ")
    keys["defence"] = "defense"
    
    print(keys)
    print(len(keys))

    peacefiles_w = [join(peace_dir_w,f) for f in listdir(peace_dir_w) if isfile(join(peace_dir_w, f)) and ".docx" in f and "~$" not in f]
    #logging.info("Peace treaties: %d"%(len(peacefiles_w)))

    commfiles_w = [join(comm_dir_w,f) for f in listdir(comm_dir_w) if isfile(join(comm_dir_w, f)) and ".docx" in f and "~$" not in f]
    # logging.info("Comm treaties: %d"%(len(commfiles_w)))
    # 
    # files_w = peacefiles_w + commfiles_w
    peacefiles_w.sort()
    commfiles_w.sort() 
    # 
    # f = open('ShortBook.csv', 'w+')
    # short_doc = csv.writer(f)
    #short_doc.writerow(["fullBook", "shortFile", "fullFile"])
    

    for f in range(len(peacefiles_w)): 
    
        logging.info("Cleaning: %s"%(peacefiles_w[f]))
        
        doc = docx.Document(peacefiles_w[f])
        #result = [p.text for p in doc.paragraphs]
        
        fname = peacefiles_w[f].split("/")[1]
    
        #doc = docx.Document("Word docs_CommNav/6_CTS_469_eng_text.docx")
        doc_n = docx.Document()
        paras = [p.text for p in doc.paragraphs]
        paras_f = [p.runs for p in doc.paragraphs]
        
        p_n = ""
        for p in paras:
            p_n = p
            #clean some symbols before replacement
            p_n = re.sub("- ", "", p_n)
            
            for k in keys:
                #in case we have "(King)" in text, next step wont recognize this
                if "(" in k or ")"in k or "(" in keys[k] or ")" in keys[k]:
                    p_n = p_n.replace(k, keys[k])
                    continue
                
                p_n = re.sub(r"(\s%s)([\s'\",.;:}\)\]])"%(k), r" %s\2"%(str(keys[k])), str(p_n))
                #p_n = re.sub(r"(\s%s)(\W)"%("Shall"), r" Shall\2", "I Shall, do y work")
            
            doc_n.add_paragraph(p_n)
    
        doc_n.save("%s/%s"%(peace_dir_w_clean, fname))
        doc_n.save("%s/%s"%(final_dir_clean, fname))
        logging.info("Done: %s"%(fname))
 
 
    for f in range(len(commfiles_w)): 
    
        logging.info("Cleaning: %s"%(commfiles_w[f]))
        doc = docx.Document(commfiles_w[f])
        #result = [p.text for p in doc.paragraphs]
        fname = commfiles_w[f].split("/")[1]
    
        #doc = docx.Document("Word docs_CommNav/6_CTS_469_eng_text.docx")
        doc_n = docx.Document()
        paras = [p.text for p in doc.paragraphs]
        paras_f = [p.runs for p in doc.paragraphs]
        
        p_n = ""
        for p in paras:
            #pdb.set_trace() 
            p_n = p
            p_n = re.sub("- ", "", p_n)
            
            for k in keys:

                if "(" in k or ")"in k or "(" in keys[k] or ")" in keys[k]:
                    p_n = p_n.replace(k, keys[k])
                    continue
                
                p_n = re.sub(r"(\s%s)([\s'\",.;:}\)\]])"%(k), r" %s\2"%(str(keys[k])), str(p_n))
                #p_n = re.sub(r"(\s%s)(\W)"%("Shall"), r" Shall\2", "I Shall, do y work")
            
            doc_n.add_paragraph(p_n)
    
        doc_n.save("%s/%s"%(comm_dir_w_clean, fname))
        doc_n.save("%s/%s"%(final_dir_clean, fname))
        logging.info("Done: %s"%(fname))
        
        
def constructYearDoc():
    f =  open('alltreaties_metadata.csv', 'r')
    corrs = csv.DictReader(f)
    keys = {}
    for c in corrs:
        try:
            #print( "****", c.split(',')[11].strip(" "))
            file = "_".join(c["oup_reference"].split(" "))
            year =  c["year"].strip(" ")
            keys[file] = int(year)
        except Exception:
            logging.info("miss %s ", c['oup_reference'].strip(" "))
            continue
            
    print(keys)
    
    f =  open('BooksPeaceComm.csv', 'w')
    fieldnames = ['first_name', 'last_name']
    cors = csv.DictWriter(f)
    
  
def test1():
    # peacefiles_w = [join(peace_dir_w,f) for f in listdir(peace_dir_w) if isfile(join(peace_dir_w, f)) and ".docx" in f]
    # for f in range(len(peacefiles_w)): 
    # 
    #     logging.info("Splitting: %s"%(peacefiles_w[f]))
    #     
    #     doc = docx.Document(peacefiles_w[f])
    #     section = doc.sections[0]
    #     header = section.header
    #     #pdb.set_trace()
    #     for paragraph in header.paragraphs:
    #         print(paragraph.text) # or whatever you have in mind
    
    
    import geograpy
    import nltk
    from geograpy import extraction
    
    text="Thyroid-associated orbitopathy (TO) is an autoimmune-\
    mediated orbital inflammation that can lead to disfigurement and blindness. \
    Multiple genetic loci have been associated with Graves' disease, but the genetic \
    basis for TO is largely unknown. British This japanese study aimed to identify loci associated with \
    TO in individuals with Graves' Mexican austrian japan Chinese disease, using a genome-wide association scan \
    (GWAS) for the first time to our knowledge in TO.Genome-wide association scan was \
    performed on pooled USA united states DNA from French britain an Australian Caucasian discovery cohort of 265 \
    participants with Graves' disease and TO (cases) and 147 patients with Graves' \
    mapisease without TO (controls)."
    
    from geotext import GeoText
    places = GeoText(text)
    
    print(places)


    e = extraction.Extractor(text=text)
    nes = nltk.ne_chunk(nltk.pos_tag(text))
    #print(nes)
    e.find_geoEntities()
    #print(e.find_geoEntities())
    country = e.places
    for s in range(len(country)):
        country[s] = stemmer.stem(country[s])
    
    # You can now access all of the places found by the Extractor
    print(country)
    #pdb.set_trace()
    print( stemmer.stem("Chinese"))
  
def test2():
    r = open("ShortFinalKeys_g.txt", "r")
    #pdb.set_trace()
    ws= r.readlines()
    extra = ['x', 'hanlfa', 'dirham', 'cabd', 'dimna', 'shafici', 'mu', 'oh', 'schanzabeh', 'culama', 'll', \
    'got', 'telemachu', 'ulyss', 'idomeneu', 'itali', 'ithaca', 'england', 'french', 'philip', 'adrastu', \
    'troy', 'thoma', 'paul', 'cicero', 'ye', 'whereof', 'loui', 'henri', 'israel', 'st', 'aforesaid'] 
    print(len(ws))
    for i in ws:
        print(len(ws))
        if i.strip("\n") in extra:
            print(i)
            ws.remove(i)
    print(len(ws))
    r.close()
    #pdb.set_trace()
    print(len(ws))
    f = open("ShortFinalKeys_g.txt", "w+")
    f.truncate(0)
    j = 0
    for w in ws:
        print (j, w)
        j += 1
      
        f.write(w)
        #r.write("\n")
            

def test3():
    # from autocorrect import Speller
    # doc = docx.Document("Word docs_Peace/1_CTS_119_eng_text.docx")
    # result = [p.text for p in doc.paragraphs]
    # 
    # spell = Speller(lang='en')
    # 
    # for j in range(15):
    #     print(spell(result[j]))

    # import jamspell
    # 
    # corrector = jamspell.TSpellCorrector()
    # corrector.LoadLangModel('en.bin')
    # text = "tended by one againft another upon this account, fhall\nbe bury'd in perpetual Oblivion.\nIII. According to this Foundation of a general and un-\nlimited Amnefty, all and every the Electors of the Sa-\ncred Roman Enmpire, the Princes and States therein inclu-\nded, the Nobility that hold immediately of the Empire,\ntheir Vaffals, Subjects, Citizens and Inhabitants, who\nupon occafion of the Troubles of Bohemia and Germany,\nor upon the account of Alliances contracted on one fide\nand another, may have fuffer'd any Prejudice or Damage\nfrom either Party, in any manner, or under any pretext\nwhatfoever, either in their Domains, Goods, Fees,\nSub-Fees, Állodials, or in their Dignities, Immunities,\nRights and Privileges, fhal be fully re-eftablifh'd on both\nfides, in the fame Štate, both as to Spirituals and Tem-\nporals, which they enjoy'd, or could of Right enjoy be-\nfore thofe Troubles, notwithftanding all the Changes\nmade to the contrary, which fhall be annul'd and remain\nvoid.\nBut as thefe and fuch like Reftitutions ought to be al\nunderftood, faving whatfoever Rights, either of Domi-\nnium directum, or Dominium utile, go along with the\nGoods which are to be reftor'd, whether Secular or Ec-\nclefiaftical, and belong to him who makes Reftitution,\nor to him to whom Reftitution is made, or to any third\nPerfon; faving alfo the Rights which lie undeternin'd ei-\nther in the Imperial Court, or in the Imperial Chamber,\n",
    # 
    # text = corrector.FixFragment(text)
    # print(text)
    sys.path.append("treatyUtil")
    import pkg_resources
    from symspellpy import SymSpell, Verbosity
    from treatyUtil import spellcheck_keep_punctuation

    sym_spell = SymSpell(max_dictionary_edit_distance=2, prefix_length=7)
    dictionary_path = pkg_resources.resource_filename(
        "symspellpy", "frequency_dictionary_en_82_765.txt")
    bigram_path = pkg_resources.resource_filename(
        "symspellpy", "frequency_bigramdictionary_en_243_342.txt")
# term_index is the column of the term and count_index is the
# column of the term frequency
    sym_spell.load_dictionary(dictionary_path, term_index=0, count_index=1)
    sym_spell.load_bigram_dictionary(bigram_path, term_index=0, count_index=2)
    

# lookup suggestions for multi-word input strings (supports compound
# splitting & merging)
    input_term1 = "tended by one againft another upon this account, fhall\nbe bury'd in perpetual Oblivion.\nIII.\
    According to this Foundation of a general and un-\nlimited Amnefty, all and every the Electors of the Sa-\ncred \
    Roman Enmpire, the Princes and States therein inclu-\nded, the Nobility that hold immediately of the Empire,\ntheir \
    Vaffals, Subjects, Citizens and Inhabitants, who\nupon occafion of the Troubles of Bohemia and Germany,\nor upon the \
    account of Alliances contracted on one fide\nand another, may have fuffer'd any Prejudice or Damage\nfrom either \
    Party, in any manner, or under any pretext\nwhatfoever, either in their Domains, Goods, Fees,\nSub-Fees, Állodials, \
    or in their Dignities, Immunities,\nRights and Privileges, fhal be fully re-eftablifh'd on both\nfides, in the fame Štate, \
    both as to Spirituals and Tem-\nporals, which they enjoy'd, or could of Right enjoy be-\nfore thofe Troubles, notwithftanding \
    all the Changes\nmade to the contrary, which fhall be annul'd and remain\nvoid.\nBut as thefe and fuch like Reftitutions \
    ought to be al\nunderftood, faving whatfoever Rights, either of Domi-\nnium directum, or Dominium utile, go along with \
    the\nGoods which are to be reftor'd, whether Secular or Ec-\nclefiaftical, and belong to him who makes Reftitution,\nor \
    to him to whom Reftitution is made, or to any third\nPerfon; faving alfo the Rights which lie undeternin'd ei-\nther in the\
    Imperial Court, or in the Imperial Chamber,\n"
    #input_term = "tended by one againft another upon this account, fhall\nbe bury'd in perpetual Oblivion.\nIII. According to this Foundation of a general and un-\nlimited "
    
    input_term = "God, and Safety of the Chriſtian World (the Electors,\nPrinces and States of the Sacred Roman Empire \
    being\npreſent, approving and conſenting) the Articles of Peace\nand Anity, whereof the Tenour follows.\n1. That \
    there be a Chriſtian, univerſal\nThe Re-efta. and perpetual Peace, and a true and ſincere\nbliſhment of Friendſhip and \
    Amity between his Sacred\nPeace and A. Imperial Majeſty, the Houſe of Austria,\nmity.\nand all his Allies and Adherents, \
    and the\nHeirs and Succeffors of each of them, chiefly the King\nof Spain, and the Electors, Princes and States of the En-\npire,\
    of the one ſide, and her Sacred Royal Majeſty,\nand the Kingdom of Sweden, her Allies and Adherents,\nand the Heirs and Succeſſors\
    of each of them, eſpecially\nthe moſt Chriſtian King, the reſpective Electors, Princes\nand States of the Empire, of the other ſide ; \
    and that this\nPeace be obſerv'd and cultivated ſincerely and ſeriouſly,\nſo that each Party may procure the Benefit, Honour and\nAdvantage \
    of one another, and thereby the Fruits of this\nPeace and Amity may be ſeen to grow up and fouriſh a-\nnew, by a ſure and reciprocal \
    maintaining of a good\nand faithful Neighbourhood between the Roman Empire\nand the Kingdom of Sweden reciprocally,\nII. That there be \
    on both ſides à perpe-\nAn Amneſty\ntua) Oblivion and Amneſty of all that has\nfrom all Hoffi- been done Since the beginning of theſe\nlity.\nTroubles, \
    in what Place or in what Man-\n"
    input_term2 = "God, and Safety of the Chriſtian World (the Electors,\nPrinces"
    input_term = re.sub("\n", " ", input_term)
    input_term = re.sub("- ", "", input_term)
    #input_term = re.sub("-", "", input_term)        
    input_term = re.sub("ſ", "s", input_term)
    
    # word_split = re.compile(r"[^\W]+", re.U)
    # suggestions = sym_spell.lookup_compound((input_term), ignore_non_words=True, max_edit_distance=2)
    # for suggestion in suggestions:
    #    print(suggestion)
    # 
    # corrected = suggestions[0].term
    # # This combined with split_phrase_by_space=True would be enough just to spell check
    # # but punctuation is lost.
    # 
    # # The spell check is already done in 'corrected'. Now we just want to keep the punctuation.
    # in_list = word_split.findall(input_term)
    # chk_list = word_split.findall(corrected)
    # print(input_term)
    # print(corrected)
    # print(in_list)
    # print(chk_list)
    # pdb.set_trace()
    # 
    # # To keep punctuation we take the original phrase and do word by word replacement
    # out_term = ""
    # outs  = input_term.split()
    # word_count = 0
    # for word in in_list:
    #     print(out_term)
    #     print(outs[word_count].lower(), word, chk_list[word_count])
    #     temp = outs[word_count].lower().replace(word, chk_list[word_count])
    #     word_count += 1
    #     out_term += temp+" "
    # 
    # print(out_term) 
    # return

# max edit distance per lookup (per single word, not per whole input string)
    #pdb.set_trace()
    #print(spellcheck_keep_punctuation(input_term))
    suggestions = sym_spell.lookup_compound((input_term), 
    transfer_casing = True,
    ignore_non_words=True, max_edit_distance=2)
# display suggestion term, edit distance, and term frequency
    #print(suggestions)
    for suggestion in suggestions:
       print(suggestion)
    

def batch_ocr_convert():
  
    if os.path.exists("ocr_result"):    
        shutil.rmtree("ocr_result")   
    os.mkdir('ocr_result')
        
    path = "/Users/mac/Desktop/test6.png"
    conv_paras = convert_pic(path)
    
    doc_n = docx.Document()
    for para in conv_paras:
        doc_n.add_paragraph(para)
        
    doc_n.save("ocr_result/test6.docx")
    

def verifyResult():
  
    ###
    # Check if short file folder and csv matches
    ###
    shortfile_t = [f for f in listdir(short_dir) if isfile(join(short_dir, f)) and ".txt" in f]
    logging.info("Short files total: %d"%(len(shortfile_t)))     
    
    f =  open('ShortBook.csv', 'r')
    starts = f.readlines()
    clean_starts = []
    
    dup = []
    for i in range(1, len(starts)):
        #ee = re.findall('truelaw|alchemy|consolationforruler', starts[z].split(',')[0].lower())
        #if len(ee)==0:
        if not starts[i].split(',')[1].strip('\n') in shortfile_t:
            print(starts[i].split(',')[1].strip('\n'))
        
        if (starts[i].split(',')[1].strip('\n') in dup):
            print(starts[i].split(',')[1].strip('\n'))
        
        dup.append(starts[i].split(',')[1].strip('\n'))
        
    for i in range(len(shortfile_t)):
        if not shortfile_t[i] in dup:
            print(shortfile_t[i])
    
    ###
    # Brutal search to see find any corner cases in replacing correction list
    ###
    f =  open('spelling_corrections.csv', 'r')
    corrs = f.readlines()
    keys = {}
    for c in corrs:
        keys[c.split(',')[0].strip(" ")] = c.split(",")[1].strip(" ")

    peacefiles_w = [join(peace_dir_w_clean,f) for f in listdir(peace_dir_w_clean) if isfile(join(peace_dir_w_clean, f)) and ".docx" in f and "~$" not in f]
    commfiles_w = [join(comm_dir_w_clean,f) for f in listdir(comm_dir_w_clean) if isfile(join(comm_dir_w_clean, f)) and ".docx" in f and "~$" not in f]
    
    files_w = peacefiles_w + commfiles_w
    files_w.sort()
    
    missingCorr = 0
    for f in range(len(files_w)): 
        doc = docx.Document(files_w[f])
        paras = [p.text for p in doc.paragraphs]
        
        for p in paras:
            for k in keys:
                if " "+k+" " in p:
                    logging.warning("Missing word: %s, in %s", k, files_w[f])
                    missingCorr = 1
    
    ret = 0
    if((len(starts)-1) != len(shortfile_t)): 
        os.system('say "Sorry! errors: short file mismatch"')
        ret = 1
    if(missingCorr == 1):
        os.system('say "Sorry! errors: missing corrections"')
        ret = 1
      
    return ret
        
    
    
if __name__ == "__main__":
    #cleanTexts()
    #preprocessWord()
    #genFinalKeys()
    #genPeaceCommDoc()
    #batch_ocr_convert()
    
    ret = verifyResult()
    if ret == 0:
        os.system('say "Congratulations! your program has finished"')
